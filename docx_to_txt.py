from werkzeug.utils import secure_filename
from flask import Flask, jsonify, request
import xml.dom.minidom
from tqdm import tqdm 
from docx_util import get_docx_text
from docx import Document
import pdfplumber
import docx2txt
import warnings
import zipfile
import spacy
import docx 
import json
import os  
import re
warnings.filterwarnings("ignore")
app= Flask(__name__)
updated_model_dir = '/home/softuvo/Garima/SpacyNER/updated_model'
print("Loading updated model from:", updated_model_dir)
nlp = spacy.load(updated_model_dir)

UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Uploads')
CONVERTED_TO_TXT = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Converted_To_Txt')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CONVERTED_TO_TXT, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['CONVERTED_TO_TXT'] = CONVERTED_TO_TXT

# ALLOWED_EXTENSIONS = {'pdf'} 
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'doc'}
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/nertagging', methods=['GET', 'POST'])
def index():
    try:
        if request.method == 'POST':
            if 'file' not in request.files:
                return jsonify({"message": 'No file provided/Check the key, spelling or any space after or in the key.'})
            file = request.files['file']
            if file.filename == '':
                return jsonify({"message": 'No file selected.'})
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            else:
                return jsonify({"message": f'{file.filename} not allowed. Please send pdf format only.'})
            
            filename, file_extension =  os.path.splitext(file.filename)
            print(f'{filename}, {file_extension}')
           
            if file_extension == ".docx":
                file_name = os.path.join(CONVERTED_TO_TXT, file.filename.replace('.docx', '.txt'))
                if os.path.isfile(filename):
                    os.remove(filename)
                text = open(os.path.join(UPLOAD_FOLDER, file.filename))
                print("@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@")
                document = zipfile.ZipFile(text)
                print(document)
                print("@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@")
                doc = Document(document)
                print(doc.paragraphs)
                print("@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@")
                for para in doc.paragraphs:
                    print(para.text)



                # data = ""
                # fullText = []
                # for para in doc.paragraphs:
                #     fullText.append(para.text)
                #     data = '\n'.join(fullText)
                # print(data) 
                



                
                # document = zipfile.ZipFile(text)

               
                # print("***********************************************************************************")
                
                
                
                # document = zipfile.ZipFile(file_name)
                # print("@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@")
                # print(document)

                # data = ZipFile.read(name, pwd=None)
                # with open(file_name, "w") as text_file:
	            #     print(text, file=text_file)
              
                # print(text)

            # else:
            #     return jsonify({"message": "Data not uploaded properly"})
    except Exception as ex:
        print(ex)
        return jsonify({"message": "Allowed Entensions are .pdf, .txt"})
                    
if __name__ == '__main__':
	app.run(port=5001, debug=True)